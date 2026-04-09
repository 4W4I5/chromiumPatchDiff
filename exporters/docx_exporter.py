from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt

_DEV_CHURN_PATH_PATTERNS: tuple[re.Pattern[str], ...] = (
    re.compile(r"(^|/)build\.(gn|gni)$", flags=re.IGNORECASE),
    re.compile(r"(^|/)(version|versions|lastchange|deps)$", flags=re.IGNORECASE),
    re.compile(r"(^|/)chrome/VERSION$", flags=re.IGNORECASE),
    re.compile(r"(^|/)tools/", flags=re.IGNORECASE),
    re.compile(r"/translations/", flags=re.IGNORECASE),
    re.compile(r"\.(xtb|grd|grdp|pak)$", flags=re.IGNORECASE),
)

_DEV_CHURN_TEXT_HINTS: tuple[str, ...] = (
    "version bump",
    "bump version",
    "update version",
    "roll ",
    "autoroll",
    "translation",
    "string update",
)

_SECURITY_DEFAULT_TERMS: tuple[str, ...] = (
    "cve-",
    "security",
    "overflow",
    "heap",
    "stack",
    "uaf",
    "use-after-free",
    "type confusion",
    "out of bounds",
    "oob",
    "memory corruption",
    "sandbox",
    "rce",
)

_CODE_FILE_EXTENSIONS: set[str] = {
    ".c",
    ".cc",
    ".cpp",
    ".cxx",
    ".h",
    ".hh",
    ".hpp",
    ".hxx",
    ".m",
    ".mm",
    ".java",
    ".js",
    ".ts",
    ".py",
    ".rs",
    ".go",
    ".swift",
    ".asm",
    ".s",
    ".patch",
}

_PATCH_APPENDIX_FILE_LIMIT = 40
_PATCH_APPENDIX_CHAR_LIMIT = 6000


def build_analysis_docx(result: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading("Chromium Patch Diff Analysis Report", level=0)
    document.add_paragraph(f"Generated at: {result.get('generated_at', '')}")

    _add_executive_summary(document, result)
    _add_effective_focus(document, result)
    _add_cve_details(document, result)
    _add_release_blog_details(document, result)
    _add_commit_evidence(document, result)
    _add_filtered_diff_summary(document, result)
    _add_patch_appendix(document, result)
    _add_warnings_and_provenance(document, result)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _add_executive_summary(document: Document, result: dict[str, Any]) -> None:
    document.add_heading("Executive Summary", level=1)

    input_payload = result.get("input", {}) if isinstance(result.get("input"), dict) else {}
    compare = result.get("compare", {}) if isinstance(result.get("compare"), dict) else {}

    component_status_counts: dict[str, int] = {}
    for component in compare.get("components", []) or []:
        if not isinstance(component, dict):
            continue
        status = str(component.get("status", "unknown") or "unknown")
        component_status_counts[status] = component_status_counts.get(status, 0) + 1

    status_summary = ", ".join(f"{key}={value}" for key, value in sorted(component_status_counts.items())) or "(none)"

    summary_rows = [
        ("Input mode", str(result.get("input_mode", ""))),
        ("CVE ID", str(input_payload.get("cve_id", ""))),
        ("Version input", str(input_payload.get("version", input_payload.get("version_hint", "")))),
        ("Patched version", str(result.get("patched_version", ""))),
        ("Unpatched version", str(result.get("unpatched_version", ""))),
        ("Platform", str(input_payload.get("platform", ""))),
        ("Components", ", ".join(input_payload.get("components", []) or [])),
        ("Compare commits", str(compare.get("total_commit_count", 0))),
        ("Compare files", str(compare.get("total_file_count", 0))),
        ("Component statuses", status_summary),
    ]

    _add_key_value_table(document, summary_rows)


def _add_effective_focus(document: Document, result: dict[str, Any]) -> None:
    document.add_heading("Effective Security Focus", level=1)

    focus = result.get("effective_focus", {}) if isinstance(result.get("effective_focus"), dict) else {}
    compare = result.get("compare", {}) if isinstance(result.get("compare"), dict) else {}
    filters = compare.get("filters", {}) if isinstance(compare.get("filters"), dict) else {}

    rows = [
        ("Mode", "Minimal" if bool(focus.get("minimal_mode")) else "Advanced"),
        ("Code scope", str(focus.get("code_scope", "changed-files-only"))),
        ("Components", ", ".join(focus.get("components", []) or []) or "(none)"),
        ("Auto keywords", ", ".join(focus.get("auto_keywords", []) or []) or "(none)"),
        ("Effective keywords", ", ".join(focus.get("keywords", []) or []) or "(none)"),
        ("Evidence tokens", ", ".join(filters.get("evidence_tokens", []) or []) or "(none)"),
    ]
    _add_key_value_table(document, rows)


def _add_cve_details(document: Document, result: dict[str, Any]) -> None:
    cve_payload = result.get("cve")
    if not isinstance(cve_payload, dict):
        cve_payload = result.get("primary_cve") if isinstance(result.get("primary_cve"), dict) else None

    if not isinstance(cve_payload, dict):
        cves = (result.get("enrichment", {}) if isinstance(result.get("enrichment"), dict) else {}).get("cves", [])
        if isinstance(cves, list) and cves:
            cve_payload = cves[0] if isinstance(cves[0], dict) else None

    document.add_heading("CVE Enrichment Details", level=1)

    if not isinstance(cve_payload, dict):
        document.add_paragraph("No CVE detail payload was available for this run.")
        return

    detail_rows = [
        ("CVE ID", str(cve_payload.get("cve_id", ""))),
        ("Source", str(cve_payload.get("source", ""))),
        ("Published", str(cve_payload.get("published", ""))),
        ("Updated", str(cve_payload.get("updated", ""))),
        ("Match reason", str(cve_payload.get("match_reason", ""))),
        ("Match confidence", str(cve_payload.get("match_confidence", ""))),
        ("NVD severity", str((cve_payload.get("nvd") or {}).get("severity", ""))),
        ("NVD CVSS", str((cve_payload.get("nvd") or {}).get("cvss_score", ""))),
    ]
    _add_key_value_table(document, detail_rows)

    if cve_payload.get("title"):
        document.add_paragraph(f"Title: {cve_payload.get('title', '')}")

    if cve_payload.get("description"):
        document.add_paragraph("Description:")
        document.add_paragraph(str(cve_payload.get("description", "")))

    references = cve_payload.get("references", []) or []
    if references:
        document.add_paragraph("References:")
        for item in references:
            document.add_paragraph(str(item), style="List Bullet")

    affected_versions = cve_payload.get("affected_versions", []) or []
    if affected_versions:
        document.add_paragraph("Affected version hints:")
        for item in affected_versions:
            document.add_paragraph(str(item), style="List Bullet")

    weaknesses = (cve_payload.get("nvd") or {}).get("weaknesses", []) or []
    if weaknesses:
        document.add_paragraph("NVD weaknesses:")
        for item in weaknesses:
            document.add_paragraph(str(item), style="List Bullet")


def _add_release_blog_details(document: Document, result: dict[str, Any]) -> None:
    document.add_heading("Chrome Releases Blog Matches", level=1)

    release_blog = result.get("release_blog") if isinstance(result.get("release_blog"), dict) else {}
    query_cve = str(release_blog.get("query_cve_id", "") or "").strip().upper()
    posts = release_blog.get("posts", []) if isinstance(release_blog.get("posts"), list) else []
    selected = release_blog.get("selected_log_range") if isinstance(release_blog.get("selected_log_range"), dict) else {}
    query_bug_ids = [str(item).strip() for item in (release_blog.get("query_cve_bug_ids", []) or []) if str(item).strip()]

    selected_range = ""
    if isinstance(selected, dict):
        base_version = str(selected.get("base_version", ""))
        head_version = str(selected.get("head_version", ""))
        if base_version and head_version:
            selected_range = f"{base_version}..{head_version}"

    summary_rows = [
        ("Query CVE", query_cve),
        ("Matched post count", str(len(posts))),
        ("Selected log range", selected_range or "(none)"),
        ("Selected post", str(selected.get("post_title", "") if isinstance(selected, dict) else "")),
        ("Mapped bug IDs for query CVE", ", ".join(query_bug_ids) if query_bug_ids else "(none)"),
    ]
    _add_key_value_table(document, summary_rows)

    if isinstance(selected, dict) and selected.get("log_url"):
        document.add_paragraph(f"Selected log URL: {selected.get('log_url', '')}")

    if not posts:
        document.add_paragraph("No Stable Desktop Chrome Releases post matched this CVE lookup.")
        return

    for post in posts:
        if not isinstance(post, dict):
            continue

        title = str(post.get("title", "")).strip() or "Chrome release post"
        document.add_heading(title, level=2)

        post_rows = [
            ("Post URL", str(post.get("url", ""))),
            ("Published", str(post.get("published", ""))),
            ("Updated", str(post.get("updated", ""))),
            ("Matched CVEs", ", ".join(post.get("matched_cves", []) or [])),
        ]
        _add_key_value_table(document, post_rows)

        log_links = post.get("log_links", []) or []
        if log_links:
            document.add_paragraph("Log links:")
            for item in log_links:
                if not isinstance(item, dict):
                    continue
                base = str(item.get("base_version", ""))
                head = str(item.get("head_version", ""))
                range_label = f"{base}..{head}" if base and head else "range-unparsed"
                document.add_paragraph(f"{range_label}: {item.get('url', '')}", style="List Bullet")


def _add_commit_evidence(document: Document, result: dict[str, Any]) -> None:
    document.add_heading("Commit Evidence", level=1)
    security_terms = _collect_security_terms(result)
    commits = _collect_commit_rows(result, security_terms=security_terms)

    if not commits:
        document.add_paragraph("No security-relevant commit evidence was found for this analysis.")
        return

    table = document.add_table(rows=1, cols=7)
    header_cells = table.rows[0].cells
    header_cells[0].text = "Component"
    header_cells[1].text = "SHA"
    header_cells[2].text = "Title"
    header_cells[3].text = "URL"
    header_cells[4].text = "Confidence"
    header_cells[5].text = "Mapped CVEs"
    header_cells[6].text = "Matched Bug IDs"

    for row in commits:
        cells = table.add_row().cells
        cells[0].text = row.get("component", "")
        cells[1].text = row.get("sha", "")
        cells[2].text = row.get("title", "")
        cells[3].text = row.get("url", "")
        cells[4].text = row.get("confidence", "")
        cells[5].text = row.get("mapped_release_cves", "")
        cells[6].text = row.get("matched_release_bug_ids", "")


def _add_filtered_diff_summary(document: Document, result: dict[str, Any]) -> None:
    document.add_heading("Compare Results (Security-Relevant)", level=1)

    compare = result.get("compare", {}) if isinstance(result.get("compare"), dict) else {}
    security_terms = _collect_security_terms(result)

    summary_rows = [
        ("Base version", str(compare.get("base_version", ""))),
        ("Head version", str(compare.get("head_version", ""))),
        ("Platform", str(compare.get("platform", ""))),
        ("Total components", str(compare.get("total_component_count", 0))),
        ("Total commits", str(compare.get("total_commit_count", 0))),
        ("Total files", str(compare.get("total_file_count", 0))),
    ]
    _add_key_value_table(document, summary_rows)

    components = compare.get("components", []) or []
    for component_result in components:
        if not isinstance(component_result, dict):
            continue

        component_name = str(component_result.get("component", "")).strip() or "unknown"
        document.add_heading(f"Component: {component_name}", level=2)

        all_files = [item for item in (component_result.get("files", []) or []) if isinstance(item, dict)]
        security_files = [item for item in all_files if _is_security_relevant_file(item, security_terms)]
        suppressed_count = max(0, len(all_files) - len(security_files))
        fallback_hint = component_result.get("fallback_version_hint", {}) if isinstance(component_result.get("fallback_version_hint"), dict) else {}
        fallback_hint_version = str(fallback_hint.get("suggested_chromium_version", "") or "").strip()
        fallback_hint_build = str(fallback_hint.get("suggested_build_number", "") or "").strip()
        fallback_hint_strategy = str(fallback_hint.get("strategy", "") or "").strip()

        component_rows = [
            ("Status", str(component_result.get("status", ""))),
            ("Repo", str(component_result.get("repo", ""))),
            ("Compare URL", str(component_result.get("compare_url", ""))),
            ("Commit count", str(component_result.get("commit_count", 0))),
            ("File count (raw)", str(component_result.get("file_count", 0))),
            ("File count (security-relevant)", str(len(security_files))),
            ("Suppressed non-security files", str(suppressed_count)),
            ("Truncated", str((component_result.get("compare_meta") or {}).get("truncated", False))),
            (
                "Fallback version/build hint",
                (
                    f"{fallback_hint_version} (build {fallback_hint_build}) via {fallback_hint_strategy}"
                    if bool(fallback_hint.get("applied")) and fallback_hint_version
                    else "(none)"
                ),
            ),
        ]
        _add_key_value_table(document, component_rows)

        if security_files:
            table = document.add_table(rows=1, cols=5)
            table.rows[0].cells[0].text = "Filename"
            table.rows[0].cells[1].text = "Status"
            table.rows[0].cells[2].text = "Additions"
            table.rows[0].cells[3].text = "Deletions"
            table.rows[0].cells[4].text = "Changes"

            for file_item in security_files:
                row = table.add_row().cells
                row[0].text = str(file_item.get("filename", ""))
                row[1].text = str(file_item.get("status", ""))
                row[2].text = str(file_item.get("additions", 0))
                row[3].text = str(file_item.get("deletions", 0))
                row[4].text = str(file_item.get("changes", 0))
        else:
            document.add_paragraph("No security-relevant files remained after suppressing obvious development churn.")


def _add_patch_appendix(document: Document, result: dict[str, Any]) -> None:
    document.add_heading("Security Patch Appendix", level=1)

    compare = result.get("compare", {}) if isinstance(result.get("compare"), dict) else {}
    components = compare.get("components", []) or []
    security_terms = _collect_security_terms(result)

    has_patch = False
    included_count = 0
    truncated_due_to_limit = False
    for component_result in components:
        if not isinstance(component_result, dict):
            continue

        component_name = str(component_result.get("component", "")).strip() or "unknown"
        files = [item for item in (component_result.get("files", []) or []) if isinstance(item, dict)]

        for file_item in files:
            if not _is_security_relevant_file(file_item, security_terms):
                continue
            if included_count >= _PATCH_APPENDIX_FILE_LIMIT:
                truncated_due_to_limit = True
                break

            patch_text = str(file_item.get("patch", "") or "")
            filename = str(file_item.get("filename", "") or "")
            if not patch_text:
                continue

            has_patch = True
            included_count += 1
            document.add_heading(f"{component_name}: {filename}", level=2)

            patch_excerpt = patch_text[:_PATCH_APPENDIX_CHAR_LIMIT]
            if len(patch_text) > _PATCH_APPENDIX_CHAR_LIMIT:
                patch_excerpt += "\n... patch excerpt truncated ..."

            paragraph = document.add_paragraph()
            run = paragraph.add_run(patch_excerpt)
            run.font.name = "Consolas"
            run.font.size = Pt(8)

        if truncated_due_to_limit:
            break

    if not has_patch:
        document.add_paragraph("No security-relevant textual patch data was returned by the compare sources for this analysis.")
    elif truncated_due_to_limit:
        document.add_paragraph(
            f"Patch appendix was limited to the first {_PATCH_APPENDIX_FILE_LIMIT} security-relevant files to keep the report concise."
        )


def _add_warnings_and_provenance(document: Document, result: dict[str, Any]) -> None:
    document.add_heading("Warnings and Data Source Provenance", level=1)

    warnings = result.get("warnings", []) or []
    if warnings:
        document.add_paragraph("Warnings:")
        for item in warnings:
            document.add_paragraph(str(item), style="List Bullet")
    else:
        document.add_paragraph("Warnings: none")

    provenance = result.get("provenance", []) or []
    if provenance:
        document.add_paragraph("Provenance:")
        for item in provenance:
            document.add_paragraph(str(item), style="List Bullet")


def _collect_commit_rows(result: dict[str, Any], *, security_terms: set[str]) -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []
    seen: set[tuple[str, str, str]] = set()

    cve = result.get("cve") if isinstance(result.get("cve"), dict) else None
    if isinstance(cve, dict):
        for commit in cve.get("commits", []) or []:
            if not isinstance(commit, dict):
                continue
            row = {
                "component": "cve-linked",
                "sha": str(commit.get("sha", "")),
                "title": str(commit.get("title", "")),
                "url": str(commit.get("url", "")),
                "confidence": str(commit.get("confidence", "")),
                "mapped_release_cves": "",
                "matched_release_bug_ids": "",
            }
            key = (row["component"], row["sha"], row["url"])
            if key in seen:
                continue
            seen.add(key)
            rows.append(row)

    compare = result.get("compare", {}) if isinstance(result.get("compare"), dict) else {}
    for component_result in compare.get("components", []) or []:
        if not isinstance(component_result, dict):
            continue
        component_name = str(component_result.get("component", "")).strip() or "unknown"
        for commit in component_result.get("commits", []) or []:
            if not isinstance(commit, dict):
                continue

            if not _is_security_relevant_commit(commit, security_terms):
                continue

            row = {
                "component": component_name,
                "sha": str(commit.get("sha", "")),
                "title": str(commit.get("title", "")),
                "url": str(commit.get("url", "")),
                "confidence": str(commit.get("confidence", "")),
                "mapped_release_cves": ", ".join(commit.get("mapped_release_cves", []) or []),
                "matched_release_bug_ids": ", ".join(commit.get("matched_release_bug_ids", []) or []),
            }
            key = (row["component"], row["sha"], row["url"])
            if key in seen:
                continue
            seen.add(key)
            rows.append(row)

    return rows


def _collect_security_terms(result: dict[str, Any]) -> set[str]:
    terms: set[str] = set(_SECURITY_DEFAULT_TERMS)

    input_payload = result.get("input", {}) if isinstance(result.get("input"), dict) else {}
    cve_id = str(input_payload.get("cve_id", "") or "").strip().lower()
    if cve_id:
        terms.add(cve_id)

    focus = result.get("effective_focus", {}) if isinstance(result.get("effective_focus"), dict) else {}
    for key in ("keywords", "auto_keywords", "manual_keywords"):
        for item in focus.get(key, []) or []:
            token = str(item or "").strip().lower()
            if token and (len(token) >= 3 or token.startswith("cve-")):
                terms.add(token)

    compare = result.get("compare", {}) if isinstance(result.get("compare"), dict) else {}
    filters = compare.get("filters", {}) if isinstance(compare.get("filters"), dict) else {}
    for key in ("keywords", "hard_keywords", "soft_keywords", "evidence_tokens"):
        for item in filters.get(key, []) or []:
            token = str(item or "").strip().lower()
            if token and (len(token) >= 3 or token.startswith("cve-")):
                terms.add(token)

    return terms


def _is_security_relevant_commit(commit: dict[str, Any], security_terms: set[str]) -> bool:
    mapped_cves = [str(item).strip() for item in (commit.get("mapped_release_cves", []) or []) if str(item).strip()]
    if mapped_cves:
        return True

    title = str(commit.get("title", "") or "")
    message = str(commit.get("message", "") or "")
    url = str(commit.get("url", "") or "")
    haystack = f"{title}\n{message}\n{url}".lower()

    if _matches_any_term(haystack, security_terms):
        return True

    if _looks_like_dev_churn_text(haystack):
        return False

    confidence = _safe_float(commit.get("confidence", 0.0))
    return confidence >= 0.8


def _is_security_relevant_file(file_item: dict[str, Any], security_terms: set[str]) -> bool:
    filename = str(file_item.get("filename", "") or "").strip()
    patch = str(file_item.get("patch", "") or "")
    if not filename:
        return False

    lowered_filename = filename.lower()
    haystack = f"{lowered_filename}\n{patch.lower()}"

    if _matches_any_term(haystack, security_terms):
        return True

    if _looks_like_dev_churn_file(lowered_filename, haystack):
        return False

    extension = ""
    if "." in lowered_filename:
        extension = "." + lowered_filename.rsplit(".", 1)[1]

    return extension in _CODE_FILE_EXTENSIONS


def _matches_any_term(haystack: str, terms: set[str]) -> bool:
    normalized_haystack = str(haystack or "").lower()
    for term in terms:
        token = str(term or "").strip().lower()
        if not token:
            continue
        if token in normalized_haystack:
            return True
    return False


def _looks_like_dev_churn_text(haystack: str) -> bool:
    normalized_haystack = str(haystack or "").lower()
    return any(hint in normalized_haystack for hint in _DEV_CHURN_TEXT_HINTS)


def _looks_like_dev_churn_file(lowered_filename: str, haystack: str) -> bool:
    filename = str(lowered_filename or "")
    if any(pattern.search(filename) for pattern in _DEV_CHURN_PATH_PATTERNS):
        return True
    return _looks_like_dev_churn_text(haystack)


def _safe_float(value: Any) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


def _add_key_value_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"

    for key, value in rows:
        row = table.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)
